"""VERO Parser — DOCX: Structure-aware parsing with heading and table preservation.

Uses python-docx to extract:
  - Paragraph text with heading styles (Heading 1 → #, Heading 2 → ##, etc.)
  - Tables → Markdown table format
  - Image count for metadata
  - Equation detection via math symbol scanning
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument

from app.parsers.contracts import ParsedDocument, table_to_markdown, has_math_symbols

logger = logging.getLogger(__name__)


async def parse_docx(filepath: str) -> dict:
    """Parse a DOCX file preserving headings, tables, and structure.

    Returns {"text": str, "metadata": dict, "parsed_doc": ParsedDocument}
    """
    filename = Path(filepath).name
    doc = DocxDocument(filepath)

    sections = []
    tables_raw = []
    image_count = 0
    has_equations = False

    # Extract paragraphs with heading awareness
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()
        if "heading 1" in style:
            sections.append(f"# {text}")
        elif "heading 2" in style:
            sections.append(f"## {text}")
        elif "heading 3" in style:
            sections.append(f"### {text}")
        elif "heading 4" in style:
            sections.append(f"#### {text}")
        elif "title" in style:
            sections.append(f"# {text}")
        else:
            sections.append(text)

        if has_math_symbols(text):
            has_equations = True

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows and len(rows) >= 2:
            tables_raw.append(rows)
            md = table_to_markdown(rows)
            if md:
                sections.append(f"\n{md}")

    # Count images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1

    markdown_text = "\n\n".join(sections)

    parsed_doc = ParsedDocument(
        source_type="docx",
        filename=filename,
        markdown_text=markdown_text,
        page_count=max(1, len(doc.paragraphs) // 30),
        has_images=image_count > 0,
        has_tables=len(tables_raw) > 0,
        has_equations=has_equations,
        is_slide_format=False,
        complexity="full_pipeline" if (image_count > 0 or tables_raw) else "text_only",
        image_descriptions=[],
        tables_raw=tables_raw,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(tables_raw),
            "image_count": image_count,
        },
    )

    logger.info(
        "DOCX parsed [%s]: %d chars, %d paragraphs, %d tables, %d images",
        filename, len(markdown_text), len(doc.paragraphs), len(tables_raw), image_count,
    )

    return {
        "text": markdown_text,
        "metadata": parsed_doc.metadata,
        "parsed_doc": parsed_doc,
    }
